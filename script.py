"""
script.py — Automated Assessment Marking Script
================================================
Usage (called by server.py automatically):
    python script.py <submission_folder_path> <student_email>

Output:
    Prints a JSON object to stdout. server.py parses this.

Requires:
    pip install python-docx

Template Notes:
    - Replace EXPECTED_TITLE, EXPECTED_BODY, TASKS, and all
      check_* functions to adapt this to a new assessment.
    - Always ensure the final JSON contains at minimum:
        { "score": int, "total": int, "grade": str, "tasks": [...] }
"""

import sys
import os
import json
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# ═══════════════════════════════════════════════════════
#  ASSESSMENT CONFIGURATION  ← Edit for each new task
# ═══════════════════════════════════════════════════════

ASSESSMENT_NAME = "Week 1 — Word Processing Skills"

EXPECTED_TITLE = "The Importance of Healthy Living"
EXPECTED_BODY  = (
    "Healthy living is essential for every person. It includes eating a balanced "
    "diet, exercising regularly, and getting enough sleep. Many young people do "
    "not pay attention to their health habits, which can lead to problems later in life"
)

# Task labels must match order of results returned by mark_document()
TASKS = [
    "Title bold",
    "Title centred",
    "Title size 18",
    "Title colour applied",
    "Spelling corrected",
    "Grammar corrected",
    "Body font Calibri",
    "Body font size 14",
    "Title correctly capitalised",
    "Double line spacing applied",
    "Justified text alignment",
    "Bullet list present",
    "Bold keywords (≥2 of 3)",
    "Italics applied to last paragraph",
    "Word replaced (essential / not important)",
    "Table inserted",
    "Header contains name",
    "Header contains date",
    "Footer contains page number",
    "Image inserted",
]

# ═══════════════════════════════════════════════════════
#  DOCUMENT HELPERS
# ═══════════════════════════════════════════════════════

def load_doc(path: str) -> Document:
    return Document(path)

def get_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)

# ═══════════════════════════════════════════════════════
#  CHECK FUNCTIONS  ← Replace / extend per assessment
# ═══════════════════════════════════════════════════════

def check_title(doc: Document) -> list[bool]:
    """Title paragraph: bold, centred, 18pt, coloured."""
    p = doc.paragraphs[0]
    return [
        any(r.bold for r in p.runs),
        p.alignment == WD_ALIGN_PARAGRAPH.CENTER,
        any(r.font.size and r.font.size.pt == 18 for r in p.runs),
        any(r.font.color and r.font.color.rgb for r in p.runs),
    ]

def check_spelling_grammar(text: str) -> list[bool]:
    """Spelling fix (important→essential) and grammatical comma clause."""
    return [
        "healthy living is important" not in text.lower(),
        ", which" in text,
    ]

def check_font(doc: Document) -> list[bool]:
    """Body runs: Calibri, 14pt."""
    runs = [r for p in doc.paragraphs[1:] for r in p.runs if r.text.strip()]
    if not runs:
        return [False, False]
    calibri = all(r.font.name and r.font.name.lower() == "calibri" for r in runs)
    size14  = all(r.font.size and r.font.size.pt == 14 for r in runs)
    return [calibri, size14]

def check_caps(doc: Document) -> list[bool]:
    """Title text matches expected capitalisation exactly."""
    return [doc.paragraphs[0].text.strip() == EXPECTED_TITLE]

def check_spacing_justify(doc: Document) -> list[bool]:
    """Body paragraphs: double-spaced and justified."""
    body = [p for p in doc.paragraphs[1:] if p.text.strip()]

    def is_justified(p):
        return p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    def is_double_spaced(p):
        fmt = p.paragraph_format
        return (
            fmt.line_spacing_rule == WD_LINE_SPACING.DOUBLE or
            (fmt.line_spacing and abs(float(fmt.line_spacing) - 2.0) < 0.15)
        )

    spacing = all(is_double_spaced(p) for p in body) if body else False
    justify = any(is_justified(p) for p in body)
    return [spacing, justify]

def check_bullets(doc: Document) -> list[bool]:
    """At least one list-style paragraph."""
    return [any(p.style.name.startswith("List") for p in doc.paragraphs)]

def check_bold_keywords(doc: Document) -> list[bool]:
    """At least 2 of 3 keywords bolded."""
    keywords = ["balanced diet", "exercise", "sleep"]
    found = sum(
        1 for p in doc.paragraphs for r in p.runs
        if r.bold and r.text.strip().lower() in keywords
    )
    return [found >= 2]

def check_italics(doc: Document) -> list[bool]:
    """Last paragraph has at least one italic run."""
    last = doc.paragraphs[-1]
    return [any(r.italic for r in last.runs)]

def check_replace(text: str) -> list[bool]:
    """'important' replaced with 'essential'."""
    return ["essential" in text.lower() and "important" not in text.lower()]

def check_table(doc: Document) -> list[bool]:
    """Document contains at least one table."""
    return [len(doc.tables) >= 1]

def check_header_footer(doc: Document) -> list[bool]:
    """Header: non-empty and contains a digit (date). Footer: page number '1'."""
    section = doc.sections[0]
    header  = section.header.paragraphs[0].text if section.header.paragraphs else ""
    footer  = section.footer.paragraphs[0].text if section.footer.paragraphs else ""
    return [
        bool(header.strip()),
        any(ch.isdigit() for ch in header),
        "1" in footer,
    ]

def check_image(doc: Document) -> list[bool]:
    """At least one inline image."""
    return [len(doc.inline_shapes) > 0]

# ═══════════════════════════════════════════════════════
#  MARK ENGINE
# ═══════════════════════════════════════════════════════

def mark_document(path: str) -> list[bool]:
    doc  = load_doc(path)
    text = get_text(doc)
    results: list[bool] = []
    results += check_title(doc)
    results += check_spelling_grammar(text)
    results += check_font(doc)
    results += check_caps(doc)
    results += check_spacing_justify(doc)
    results += check_bullets(doc)
    results += check_bold_keywords(doc)
    results += check_italics(doc)
    results += check_replace(text)
    results += check_table(doc)
    results += check_header_footer(doc)
    results += check_image(doc)
    return results

# ═══════════════════════════════════════════════════════
#  GRADING
# ═══════════════════════════════════════════════════════

def assign_grade(pct: float) -> str:
    if pct >= 80: return "A"
    if pct >= 70: return "B"
    if pct >= 60: return "C"
    if pct >= 50: return "D"
    return "F"

# ═══════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════

def main():
    if len(sys.argv) < 3:
        print(json.dumps({"error": "Usage: script.py <folder_path> <email>"}))
        sys.exit(1)

    folder_path = sys.argv[1]
    email       = sys.argv[2].strip().lower()

    # Locate .docx in the submission folder
    try:
        docx_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".docx")]
    except FileNotFoundError:
        print(json.dumps({"error": f"Submission folder not found: {folder_path}"}))
        sys.exit(1)

    if not docx_files:
        print(json.dumps({"error": "No .docx file found in submission folder."}))
        sys.exit(1)

    docx_path = os.path.join(folder_path, docx_files[0])

    try:
        results = mark_document(docx_path)
    except Exception as exc:
        print(json.dumps({"error": f"Marking error: {str(exc)}"}))
        sys.exit(1)

    score = int(sum(results))
    total = len(TASKS)
    pct   = round(100 * score / total, 1) if total else 0

    task_breakdown = [
        {"id": i + 1, "name": name, "passed": bool(result)}
        for i, (name, result) in enumerate(zip(TASKS, results))
    ]

    output = {
        "assessment": ASSESSMENT_NAME,
        "email":      email,
        "filename":   docx_files[0],
        "score":      score,
        "total":      total,
        "percent":    pct,
        "grade":      assign_grade(pct),
        "marked_at":  datetime.now().isoformat(),
        "tasks":      task_breakdown,
    }

    # Must be the only stdout output — server.py parses this as JSON
    print(json.dumps(output))


if __name__ == "__main__":
    main()
